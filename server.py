from fastmcp import FastMCP
import httpx
import base64
from io import BytesIO
from docx import Document
import openpyxl
import pdfplumber

mcp = FastMCP("Grist Fetcher")


@mcp.tool
async def fetch_grist_records(
    url: str,
    bearer_token: str,
    limit: int = 5
) -> dict:
    """
    Appelle une API Grist et retourne les records aplatis.

    Args:
        url: URL complete de la table Grist (sans ?limit)
        bearer_token: Token Bearer pour l authentification
        limit: Nombre de records a recuperer (defaut 5)
    """
    full_url = f"{url}?limit={limit}"

    async with httpx.AsyncClient() as client:
        response = await client.get(
            full_url,
            headers={"Authorization": f"Bearer {bearer_token}"}
        )
        response.raise_for_status()
        data = response.json()

    records = data.get("records", [])
    if not records:
        return {"records": [], "csv": "", "columns": []}

    flat_records = []
    for record in records:
        row = {"id": record.get("id")}
        fields = record.get("fields", {})
        for key, value in fields.items():
            clean_key = key.replace(" ", "_").replace("-", "_")
            if isinstance(value, list) and len(value) > 0 and value[0] == "L":
                row[clean_key] = "|".join(str(v) for v in value[1:])
            elif value is None:
                row[clean_key] = ""
            else:
                row[clean_key] = value
        flat_records.append(row)

    columns = list(flat_records[0].keys()) if flat_records else []
    csv_lines = [",".join(columns)]
    for row in flat_records:
        csv_lines.append(",".join(
            f'"{str(row.get(col, "")).replace(chr(34), chr(39))}"'
            for col in columns
        ))

    return {
        "records": flat_records,
        "csv": "\n".join(csv_lines),
        "columns": columns,
        "total_fetched": len(flat_records)
    }


@mcp.tool
async def update_huwise_dataset(
    grist_url: str,
    grist_token: str,
    huwise_dataset_uid: str,
    huwise_token: str,
    huwise_domain: str = "https://poc-gps.trial.opendatasoft.com",
    max_records: int = 100
) -> dict:
    """
    Recupere les records Grist via pagination et met a jour le dataset Huwise.

    Args:
        grist_url: URL complete de la table Grist (sans ?limit ni ?offset)
        grist_token: Bearer token Grist
        huwise_dataset_uid: UID du dataset Huwise (ex: da_xxxxx)
        huwise_token: Token d authentification Huwise
        huwise_domain: Domaine du portail Huwise
        max_records: Nombre maximum de records a charger (defaut 100)
    """
    all_records = []
    offset = 0
    limit = 20

    async with httpx.AsyncClient(timeout=60) as client:
        while len(all_records) < max_records:
            remaining = max_records - len(all_records)
            page_limit = min(limit, remaining)

            response = await client.get(
                f"{grist_url}?limit={page_limit}&offset={offset}",
                headers={"Authorization": f"Bearer {grist_token}"}
            )
            response.raise_for_status()
            data = response.json()

            records = data.get("records", [])
            if not records:
                break

            for record in records:
                row = {"id": record.get("id")}
                fields = record.get("fields", {})
                for key, value in fields.items():
                    clean_key = key.replace(" ", "_").replace("-", "_")
                    if isinstance(value, list) and len(value) > 0 and value[0] == "L":
                        row[clean_key] = "|".join(str(v) for v in value[1:])
                    elif value is None:
                        row[clean_key] = ""
                    else:
                        row[clean_key] = value
                all_records.append(row)

            if len(records) < page_limit:
                break

            offset += page_limit

    if not all_records:
        return {"success": False, "error": "Aucun record recupere depuis Grist"}

    columns = list(all_records[0].keys())
    csv_lines = [",".join(columns)]
    for row in all_records:
        csv_lines.append(",".join(
            f'"{str(row.get(col, "")).replace(chr(34), chr(39))}"'
            for col in columns
        ))
    csv_content = "\n".join(csv_lines)
    filename = f"{huwise_dataset_uid}-full.csv"

    async with httpx.AsyncClient(timeout=60) as client:
        upload_response = await client.post(
            f"{huwise_domain}/api/management/v2/datasets/{huwise_dataset_uid}/files/",
            headers={"Authorization": f"Apikey {huwise_token}"},
            files={"file": (filename, csv_content.encode("utf-8"), "text/csv")}
        )
        upload_response.raise_for_status()
        file_data = upload_response.json()
        file_uid = file_data.get("uid") or file_data.get("file_uid")

        resource_response = await client.post(
            f"{huwise_domain}/api/management/v2/datasets/{huwise_dataset_uid}/resources/",
            headers={"Authorization": f"Apikey {huwise_token}"},
            json={
                "type": "uploaded_file",
                "title": filename,
                "datasource": {
                    "type": "uploaded_file",
                    "file": {"uid": file_uid}
                }
            }
        )
        resource_response.raise_for_status()

    return {
        "success": True,
        "total_records": len(all_records),
        "columns_count": len(columns),
        "filename": filename
    }


@mcp.tool
async def extract_document(
    url: str,
    api_key: str = ""
) -> dict:
    """
    Extrait le contenu d un document depuis une URL Huwise.
    Supporte PDF (base64 pour Claude Vision), DOCX (texte natif),
    XLSX (CSV natif), et images PNG/JPEG/WEBP (base64 pour Claude Vision).

    Args:
        url: URL de l asset Huwise
        api_key: Cle API Huwise si document protege (optionnel)
    """
    headers = {}
    if api_key:
        headers["Authorization"] = f"Apikey {api_key}"
        
    # Simuler un navigateur pour éviter les blocages 403
    headers["User-Agent"] = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.get(url, headers=headers, follow_redirects=True)
        response.raise_for_status()

    content_type = response.headers.get("content-type", "").lower()
    content = response.content
    filename = url.split("/")[-1].lower().split("?")[0]

    # DOCX → texte brut natif
    if "docx" in filename or "wordprocessingml" in content_type:
        doc = Document(BytesIO(content))

        text = "\n".join([
            para.text for para in doc.paragraphs
            if para.text.strip()
        ])

        tables_csv = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                header = ",".join(f'"{c}"' for c in rows[0])
                lines = [header]
                for row in rows[1:]:
                    lines.append(",".join(f'"{c}"' for c in row))
                tables_csv.append(f"# Tableau {i+1}\n" + "\n".join(lines))

        return {
            "type": "docx",
            "text": text,
            "tables_csv": "\n\n".join(tables_csv),
            "tables_count": len(doc.tables),
            "message": f"Texte extrait nativement ({len(text)} caractères, {len(doc.tables)} tableau(x))"
        }

    # XLSX → CSV natif (toutes les feuilles)
    elif "xlsx" in filename or "spreadsheetml" in content_type:
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        sheets = {}

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Ignore les lignes entièrement vides
                if any(cell is not None for cell in row):
                    rows.append([
                        str(cell) if cell is not None else ""
                        for cell in row
                    ])

            if rows:
                csv_lines = []
                for row in rows:
                    csv_lines.append(",".join(
                        f'"{str(cell).replace(chr(34), chr(39))}"'
                        for cell in row
                    ))
                sheets[sheet_name] = "\n".join(csv_lines)

        total_sheets = len(sheets)
        total_rows = sum(
            len(csv.split("\n")) - 1
            for csv in sheets.values()
        )

        return {
            "type": "xlsx",
            "sheets": sheets,
            "sheets_count": total_sheets,
            "total_rows": total_rows,
            "message": f"XLSX extrait nativement ({total_sheets} feuille(s), ~{total_rows} lignes de données)"
        }

    # PDF → base64 pour Claude Vision
        # PDF → extraction texte native, fallback base64 pour Claude Vision
    elif "pdf" in filename or "pdf" in content_type:
        size_kb = len(content) // 1024
        try:
            with pdfplumber.open(BytesIO(content)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text)
                
                if pages_text:
                    full_text = "\n\n".join(pages_text)
                    return {
                        "type": "pdf",
                        "text": full_text,
                        "pages": len(pdf.pages),
                        "size_kb": size_kb,
                        "extraction_method": "text",
                        "message": f"PDF extrait nativement ({len(pdf.pages)} pages, {len(full_text)} caractères)"
                    }
        except Exception:
            pass

        # Fallback : base64 pour Claude Vision (PDF scanné ou illisible)
        b64 = base64.standard_b64encode(content).decode("utf-8")
        return {
            "type": "pdf",
            "base64": b64,
            "media_type": "application/pdf",
            "size_kb": size_kb,
            "extraction_method": "vision",
            "message": f"PDF scanné — envoyé à Claude Vision ({size_kb} KB)"
        }

    # Images → base64 pour Claude Vision
    elif any(ext in filename for ext in ["png", "jpg", "jpeg", "webp", "gif"]):
        media_type = content_type if content_type else "image/jpeg"
        b64 = base64.standard_b64encode(content).decode("utf-8")
        size_kb = len(content) // 1024
        return {
            "type": "image",
            "base64": b64,
            "media_type": media_type,
            "size_kb": size_kb,
            "message": f"Image prête pour Claude Vision ({size_kb} KB)"
        }

    else:
        return {
            "type": "unknown",
            "error": f"Format non supporté : {content_type} / {filename}",
            "message": "Formats supportés : PDF, DOCX, XLSX, PNG, JPEG, WEBP"
        }


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 10000))
    mcp.run(transport="http", host="0.0.0.0", port=port)
